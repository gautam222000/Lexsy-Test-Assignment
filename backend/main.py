from fastapi import FastAPI, UploadFile, File, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import List, Dict, Optional
import os
import docx
import tempfile
import json
import traceback
import uuid
from dotenv import load_dotenv
from gpt_service import GPTService

# Load environment variables from .env file (if it exists)
load_dotenv()

# Constants
DOCX_EXTENSION = '.docx'
SESSION_NOT_FOUND = "Session not found"

# Placeholder formats to try when replacing
def get_placeholder_formats(placeholder: str) -> List[str]:
    """Get all possible formats for a placeholder"""
    return [
        f"[{placeholder}]",
        f"{{{{{placeholder}}}}}",
        f"{{{{{{{placeholder}}}}}}}",
        f"{{{{{placeholder}}}}}",
        f"<{placeholder}>",
        placeholder  # Direct match
    ]

app = FastAPI()

# CORS middleware
# Get allowed origins from environment or use defaults
allowed_origins = os.getenv("ALLOWED_ORIGINS", "http://localhost:3000,http://localhost:5173").split(",")
# Add Render frontend URL if provided
render_frontend_url = os.getenv("RENDER_FRONTEND_URL")
if render_frontend_url:
    allowed_origins.append(render_frontend_url)

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize OpenAI client
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    raise ValueError("OPENAI_API_KEY environment variable is not set. Please set it in .env file or as an environment variable.")

gpt_service = GPTService(api_key=api_key)

# In-memory storage (no database needed)
document_store: Dict[str, Dict] = {}

class QuestionRequest(BaseModel):
    session_id: str
    message: Optional[str] = None

class QuestionResponse(BaseModel):
    message: Optional[str]
    is_complete: bool
    session_id: str
    placeholders_filled: List[str] = []

@app.get("/")
async def root():
    """Health check endpoint"""
    return {
        "status": "ok",
        "message": "Legal Document Filler API is running",
        "environment": "production" if os.getenv("RENDER") else "development"
    }

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy"}

@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload a document and create an assistant to analyze it"""
    import time
    start_time = time.time()
    print(f"[DEBUG] Upload started at {start_time}")
    
    if not file.filename.endswith(DOCX_EXTENSION):
        raise HTTPException(status_code=400, detail=f"Only {DOCX_EXTENSION} files are supported")
    
    # Read file content
    content = await file.read()
    
    # Save to temp file
    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=DOCX_EXTENSION)
    tmp_file.write(content)
    tmp_file.close()
    tmp_path = tmp_file.name
    
    try:
        # Upload the actual DOCX file to OpenAI
        file_id = gpt_service.upload_file(tmp_path)
        
        # Create assistant (GPT will access the file via file_search)
        assistant_id = gpt_service.create_assistant()
        
        # Create conversation thread
        thread_id = gpt_service.create_thread()
        
        # Send initial message with file attachment to analyze
        initial_message = """Please carefully analyze the attached document. 

IMPORTANT INSTRUCTIONS:
1. Read the ENTIRE document to understand its context, purpose, and structure
2. Identify ALL placeholders that need to be filled (look for patterns like [NAME], {{DATE}}, <VALUE>, etc.)
3. Understand what each placeholder means in the context of this specific document
4. After identifying placeholders, you MUST ask the user SPECIFIC, RELEVANT questions based on the document's context
5. Do NOT provide replacements or complete the document yet - you must wait for the user to answer your questions first
6. Make your questions relevant to the document's purpose and context"""
        gpt_service.send_message(thread_id, initial_message, file_ids=[file_id])
        
        # Run assistant
        print(f"[DEBUG] Running assistant, elapsed: {time.time() - start_time:.2f}s")
        run_id = gpt_service.run_assistant(thread_id, assistant_id)
        
        # Wait for response (with timeout handling)
        print(f"[DEBUG] Starting wait_for_run, elapsed: {time.time() - start_time:.2f}s")
        run_result = gpt_service.wait_for_run(thread_id, run_id, timeout=60)
        print(f"[DEBUG] wait_for_run completed, elapsed: {time.time() - start_time:.2f}s, status: {run_result.get('status')}")
        
        if run_result["status"] == "timeout":
            print(f"[DEBUG] OpenAI call timed out after {time.time() - start_time:.2f}s")
            raise HTTPException(
                status_code=504,
                detail="The document analysis is taking longer than expected. Please try again with a smaller document or wait a moment."
            )
        elif run_result["status"] == "rate_limited":
            error_info = run_result.get('error', {})
            wait_time = error_info.get('wait_time', 5)
            error_msg = error_info.get('message', 'Rate limit exceeded')
            raise HTTPException(
                status_code=429,
                detail=f"Rate limit exceeded. Please wait {wait_time:.1f} seconds and try again. Error: {error_msg}"
            )
        elif run_result["status"] != "completed":
            error_detail = run_result.get('error', {})
            if isinstance(error_detail, dict):
                error_msg = error_detail.get('message', 'Unknown error')
            else:
                error_msg = str(error_detail)
            raise HTTPException(status_code=500, detail=f"Failed to analyze document: {error_msg}")
        
        # Get assistant's response
        assistant_message = gpt_service.get_latest_assistant_message(thread_id)
        
        if not assistant_message:
            raise HTTPException(status_code=500, detail="No response from assistant")
        
        # Extract placeholders from initial analysis (if available)
        placeholders = []
        try:
            # Try to extract placeholders from assistant's response
            cleaned = gpt_service._clean_json_response(assistant_message)
            analysis = json.loads(cleaned)
            if "placeholders" in analysis:
                placeholders = analysis["placeholders"]
        except Exception:
            # If not in JSON format, GPT will handle it during conversation
            pass
        
        # Create session
        session_id = str(uuid.uuid4())
        document_store[session_id] = {
            "file_path": tmp_path,
            "assistant_id": assistant_id,
            "thread_id": thread_id,
            "placeholders": placeholders,
            "replacements": {},
            "is_complete": False
        }
        
        print(f"[DEBUG] Upload completed successfully, total elapsed: {time.time() - start_time:.2f}s")
        return {
            "session_id": session_id,
            "placeholders": placeholders,
            "message": assistant_message
        }
    
    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except Exception as e:
        print(f"[ERROR] Upload failed after {time.time() - start_time:.2f}s: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")

@app.post("/ask-question", response_model=QuestionResponse)
async def ask_question(request: QuestionRequest):
    """Have a conversation with GPT assistant to fill in placeholders"""
    try:
        session_id = request.session_id
        
        if session_id not in document_store:
            raise HTTPException(status_code=404, detail=SESSION_NOT_FOUND)
        
        session_data = document_store[session_id]
        thread_id = session_data["thread_id"]
        assistant_id = session_data["assistant_id"]
        
        # Send user message if provided
        if request.message:
            gpt_service.send_message(thread_id, request.message)
        
        # Run assistant
        run_id = gpt_service.run_assistant(thread_id, assistant_id)
        
        # Wait for response
        run_result = gpt_service.wait_for_run(thread_id, run_id)
        
        if run_result["status"] == "rate_limited":
            error_info = run_result.get('error', {})
            wait_time = error_info.get('wait_time', 5)
            error_msg = error_info.get('message', 'Rate limit exceeded')
            raise HTTPException(
                status_code=429,
                detail=f"Rate limit exceeded. Please wait {wait_time:.1f} seconds and try again. Error: {error_msg}"
            )
        elif run_result["status"] != "completed":
            error_detail = run_result.get('error', {})
            if isinstance(error_detail, dict):
                error_msg = error_detail.get('message', 'Unknown error')
            else:
                error_msg = str(error_detail)
            raise HTTPException(status_code=500, detail=f"Assistant error: {error_msg}")
        
        # Get latest assistant message
        assistant_message = gpt_service.get_latest_assistant_message(thread_id)
        
        if not assistant_message:
            raise HTTPException(status_code=500, detail="No response from assistant")
        
        # Get all messages to check if user has responded
        all_messages = gpt_service.get_messages(thread_id)
        user_message_count = sum(1 for msg in all_messages if msg["role"] == "user")
        
        # Check if conversation is complete
        # IMPORTANT: Only mark as complete if user has sent at least one message (avoid false positives on first response)
        is_complete = False
        if user_message_count > 0:
            is_complete = any(phrase in assistant_message.lower() for phrase in [
                "i have all the information needed",
                "let me complete the document now",
                "i have enough information",
                "complete the document"
            ])
        
        # Try to extract replacement mapping if complete
        replacements = {}
        if is_complete:
            mapping = gpt_service.extract_replacement_mapping(thread_id)
            if mapping:
                replacements = mapping.get("replacements", {})
                session_data["replacements"] = replacements
                if "placeholders" in mapping:
                    session_data["placeholders"] = mapping["placeholders"]
            session_data["is_complete"] = True
            # Replace the JSON response with a friendly message
            assistant_message = "Perfect! I have all the information I need. Generating your completed document now..."
        
        return QuestionResponse(
            message=assistant_message,
            is_complete=is_complete,
            session_id=session_id,
            placeholders_filled=list(replacements.keys())
        )
    except HTTPException:
        raise
    except Exception as e:
        error_traceback = traceback.format_exc()
        raise HTTPException(status_code=500, detail=f"Error in conversation: {str(e)}\n\nTraceback:\n{error_traceback}")

def replace_text_in_paragraph(paragraph, placeholder, replacement):
    """Replace placeholder text in a paragraph while preserving formatting"""
    if placeholder in paragraph.text:
        runs = paragraph.runs
        full_text = ''.join([run.text for run in runs])
        
        if placeholder in full_text:
            paragraph.clear()
            parts = full_text.split(placeholder)
            
            # Add text before placeholder with original formatting
            if parts[0]:
                if runs:
                    run = paragraph.add_run(parts[0])
                    if runs[0].bold:
                        run.bold = True
                    if runs[0].italic:
                        run.italic = True
                    if runs[0].font.size:
                        run.font.size = runs[0].font.size
                    if runs[0].font.name:
                        run.font.name = runs[0].font.name
                else:
                    paragraph.add_run(parts[0])
            
            # Add replacement text
            replacement_run = paragraph.add_run(str(replacement))
            if runs:
                for run in runs:
                    if placeholder in run.text:
                        if run.bold:
                            replacement_run.bold = True
                        if run.italic:
                            replacement_run.italic = True
                        if run.font.size:
                            replacement_run.font.size = run.font.size
                        if run.font.name:
                            replacement_run.font.name = run.font.name
                        break
            
            # Add remaining text after placeholder
            if len(parts) > 1 and parts[1]:
                if runs:
                    last_run = paragraph.add_run(parts[1])
                    if runs[-1].bold:
                        last_run.bold = True
                    if runs[-1].italic:
                        last_run.italic = True
                    if runs[-1].font.size:
                        last_run.font.size = runs[-1].font.size
                    if runs[-1].font.name:
                        last_run.font.name = runs[-1].font.name
                else:
                    paragraph.add_run(parts[1])

@app.post("/complete-document")
async def complete_document(session_id: str = Query(...), force: bool = Query(False)):
    """Complete the document by replacing placeholders ourselves
    
    Args:
        session_id: Session ID
        force: If True, complete even with partial replacements
    """
    if session_id not in document_store:
        raise HTTPException(status_code=404, detail=SESSION_NOT_FOUND)
    
    session_data = document_store[session_id]
    original_docx_path = session_data["file_path"]
    thread_id = session_data.get("thread_id")
    
    # First check if replacements are already stored
    replacements = session_data.get("replacements", {})
    placeholders_list = session_data.get("placeholders") or []
    
    # If not stored, try to extract from conversation
    if (not replacements or not placeholders_list) and thread_id:
        print(f"[DEBUG] No replacements in session_data, trying to extract from thread {thread_id}")
        mapping = gpt_service.extract_replacement_mapping(thread_id)
        if mapping:
            replacements = mapping.get("replacements", {})
            placeholders_list = mapping.get("placeholders", [])
            session_data["replacements"] = replacements
            session_data["placeholders"] = placeholders_list
        else:
            # If force=True and no replacements found, ask GPT to provide what it has
            if force:
                print("[DEBUG] Force mode: Requesting GPT to provide partial replacements...")
                assistant_id = session_data.get("assistant_id")
                if assistant_id:
                    mapping = gpt_service.request_partial_replacements(thread_id, assistant_id)
                    if mapping:
                        replacements = mapping.get("replacements", {})
                        placeholders_list = mapping.get("placeholders", [])
                        session_data["replacements"] = replacements
                        session_data["placeholders"] = placeholders_list
                        print(f"[DEBUG] Got partial replacements: {list(replacements.keys())}")
                    else:
                        print("[DEBUG] GPT did not provide replacements")
                        replacements = {}
                else:
                    replacements = {}
            else:
                # Debug: Print all messages to see what GPT said
                all_messages = gpt_service.get_messages(thread_id)
                print(f"[DEBUG] Failed to extract mapping. Total messages: {len(all_messages)}")
                for i, msg in enumerate(all_messages):
                    print(f"[DEBUG] Message {i}: role={msg['role']}, content={msg['content'][:200]}")
                
                latest_message = gpt_service.get_latest_assistant_message(thread_id)
                error_detail = latest_message[:200] if latest_message else 'No messages'
                raise HTTPException(status_code=400, detail=f"No replacements found. Please complete the conversation first. Last assistant message: {error_detail}")
    
    if not replacements and not force:
        raise HTTPException(status_code=400, detail="No replacements found. Please complete the conversation first.")
    
    print(f"[DEBUG] Using replacements: {replacements}")
    print(f"[DEBUG] Replacement keys: {list(replacements.keys())}")
    
    # Build mapping from semantic_name to literal text
    semantic_to_literal = {}
    if placeholders_list:
        for placeholder in placeholders_list:
            semantic_name = placeholder.get("semantic_name") or placeholder.get("name")  # Support both formats for backward compatibility
            literal = placeholder.get("literal")
            if semantic_name and literal:
                semantic_to_literal[semantic_name] = literal
                print(f"[DEBUG] Mapped semantic_name '{semantic_name}' -> literal '{literal}'")
    else:
        print("[WARNING] No placeholders list found, falling back to format-based replacement")
    
    # Load the original document
    doc = docx.Document(original_docx_path)
    
    # Debug: Print document text to see what we're working with
    print(f"[DEBUG] Document paragraphs: {len(doc.paragraphs)}")
    sample_text = []
    for i, para in enumerate(doc.paragraphs[:10]):  # Print first 10 paragraphs
        if para.text.strip():
            sample_text.append(para.text[:200])
            print(f"[DEBUG] Para {i}: {para.text[:200]}")
    
    # Replace placeholders in paragraphs
    replacements_applied = 0
    for paragraph in doc.paragraphs:
        for semantic_name, replacement in replacements.items():
            # Try to get literal text from placeholders mapping
            literal_text = semantic_to_literal.get(semantic_name)
            
            if literal_text:
                # Use literal text directly
                if literal_text in paragraph.text:
                    print(f"[DEBUG] Found literal placeholder '{literal_text}' in paragraph: {paragraph.text[:50]}...")
                    replace_text_in_paragraph(paragraph, literal_text, replacement)
                    replacements_applied += 1
            else:
                # Fallback to format-based replacement (for backward compatibility)
                print(f"[DEBUG] No literal found for '{semantic_name}', using format-based replacement")
                formats_to_try = get_placeholder_formats(semantic_name)
                for fmt in formats_to_try:
                    if fmt in paragraph.text:
                        print(f"[DEBUG] Found placeholder '{fmt}' in paragraph: {paragraph.text[:50]}...")
                        replace_text_in_paragraph(paragraph, fmt, replacement)
                        replacements_applied += 1
                        break
    
    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for semantic_name, replacement in replacements.items():
                        # Try to get literal text from placeholders mapping
                        literal_text = semantic_to_literal.get(semantic_name)
                        
                        if literal_text:
                            # Use literal text directly
                            if literal_text in paragraph.text:
                                print(f"[DEBUG] Found literal placeholder '{literal_text}' in table cell")
                                replace_text_in_paragraph(paragraph, literal_text, replacement)
                                replacements_applied += 1
                        else:
                            # Fallback to format-based replacement (for backward compatibility)
                            formats_to_try = get_placeholder_formats(semantic_name)
                            for fmt in formats_to_try:
                                if fmt in paragraph.text:
                                    print(f"[DEBUG] Found placeholder '{fmt}' in table cell")
                                    replace_text_in_paragraph(paragraph, fmt, replacement)
                                    replacements_applied += 1
                                    break
    
    print(f"[DEBUG] Total replacements applied: {replacements_applied}")
    
    if replacements_applied == 0:
        print("[WARNING] No replacements were applied! Check if placeholder formats match.")
        print(f"[DEBUG] Available replacement keys: {list(replacements.keys())}")
        print(f"[DEBUG] Sample document text: {sample_text[:3]}")
    
    # Save the completed document
    completed_docx_path = original_docx_path.replace(DOCX_EXTENSION, '_completed.docx')
    doc.save(completed_docx_path)
    
    # Extract text for preview
    completed_text = '\n'.join([para.text for para in doc.paragraphs])
    
    # Store completed document info
    session_data["docx_path"] = completed_docx_path
    session_data["completed_text"] = completed_text
    session_data["replacements"] = replacements
    
    print(f"[DEBUG] Document completed successfully: {completed_docx_path}")
    
    return {
        "session_id": session_id,
        "completed_text": completed_text,
        "download_url": f"/download/{session_id}",
        "replacements": replacements  # Include replacements for frontend highlighting
    }

@app.get("/document/{session_id}")
async def get_document_file(session_id: str):
    """Get the original DOCX file for preview"""
    if session_id not in document_store:
        raise HTTPException(status_code=404, detail=SESSION_NOT_FOUND)
    
    session_data = document_store[session_id]
    file_path = session_data["file_path"]
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="Document file not found")
    
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="document.docx"
    )

@app.get("/download/{session_id}")
async def download_docx(session_id: str):
    """Download the completed DOCX file"""
    if session_id not in document_store:
        raise HTTPException(status_code=404, detail=SESSION_NOT_FOUND)
    
    session_data = document_store[session_id]
    if "docx_path" not in session_data:
        raise HTTPException(status_code=404, detail="DOCX not generated yet")
    
    docx_path = session_data["docx_path"]
    if not os.path.exists(docx_path):
        raise HTTPException(status_code=404, detail="DOCX file not found")
    
    return FileResponse(
        docx_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="completed_document.docx"
    )

if __name__ == "__main__":
    import uvicorn
    import sys
    
    dev_mode = "--dev" in sys.argv or "-d" in sys.argv
    
    if dev_mode:
        uvicorn.run(
            "main:app",
            host="0.0.0.0",
            port=8000,
            reload=True,
            reload_dirs=["."],
            log_level="info"
        )
    else:
        uvicorn.run(app, host="0.0.0.0", port=8000)
